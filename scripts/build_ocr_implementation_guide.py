from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\hp\StudioProjects\MedAssist\docs\mindheal_ocr_implementation_guide.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(21, 96, 130)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        shade_cell(header_cells[index], "D9EAF7")
        if widths:
            header_cells[index].width = Inches(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Inches(widths[index])
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MindHeal OCR Implementation Guide")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(21, 96, 130)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Production prescription scanning pipeline for technical and non-technical readers"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Project: MindHeal / MedAssist | Stack: Flutter, GetX, Firebase, Google Vision, LLM parser")
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)

    add_heading(doc, "1. Plain-English Overview", 1)
    add_body(
        doc,
        "The OCR feature lets a patient take a photo of a prescription, upload it securely, read the text using Google Cloud Vision, and turn that text into structured medicine records. Those medicine records use the same Medication model that the rest of the app already uses, so saving prescriptions, scheduling reminders, and sending notifications continue to work without changes.",
    )
    add_numbered(
        doc,
        [
            "The user takes a prescription photo or chooses one from the gallery.",
            "The app uploads the image to Firebase Storage.",
            "Google Cloud Vision reads the text from the uploaded image.",
            "An LLM extracts medicines, dosage, frequency, duration, and instructions from the raw text.",
            "The controller shows a review screen where the user can save or manually adjust medicine information.",
            "When saved, the existing prescription and reminder system creates future medication doses from the extracted medicine list.",
        ],
    )

    add_heading(doc, "2. Technical Architecture", 1)
    add_table(
        doc,
        ["Layer", "Responsibility", "Main Files"],
        [
            [
                "Controller",
                "Orchestrates capture, upload, OCR, parse, state update, navigation, and user-facing errors.",
                "lib/features/prescription/controller/prescription_reader_controller.dart",
            ],
            [
                "OCR service",
                "Uploads image to Firebase Storage and calls Google Vision DOCUMENT_TEXT_DETECTION.",
                "lib/services/ocr/ocr_service.dart",
            ],
            [
                "Parser service",
                "Sends raw OCR text to the selected LLM, validates JSON, and maps output into Medication objects.",
                "lib/services/ocr/prescription_parser_service.dart",
            ],
            [
                "Environment config",
                "Loads API keys from .env using flutter_dotenv.",
                "pubspec.yaml, lib/main.dart, .env",
            ],
            [
                "Existing downstream system",
                "Saves prescription and creates reminder doses from the Medication list.",
                "PrescriptionRepository, MedicationDoseRepository",
            ],
        ],
        widths=[1.4, 3.4, 2.2],
    )

    add_heading(doc, "3. Runtime Flow", 1)
    add_code(
        doc,
        "ImagePicker -> Firebase Storage -> Google Vision DOCUMENT_TEXT_DETECTION -> raw text -> LLM JSON parser -> List<Medication> -> Prescription review -> Firestore save -> MedicationDoses -> FCM reminders",
    )
    add_body(
        doc,
        "The controller does not call Vision or the LLM directly. It only calls OcrService.scanPrescription(...). This keeps API code isolated and makes future replacement easier.",
    )

    add_heading(doc, "4. Firebase Storage Upload", 1)
    add_body(
        doc,
        "Prescription images are uploaded under a user-scoped path. This makes storage easier to inspect and avoids mixing all prescription images into a single flat folder.",
    )
    add_code(doc, "prescriptions/{userId}/{timestamp}.jpg")
    add_body(
        doc,
        "After upload, Firebase Storage returns a download URL. That URL is sent directly to Google Cloud Vision in the imageUri field.",
    )

    add_heading(doc, "5. Google Vision OCR Request", 1)
    add_body(
        doc,
        "The pipeline uses DOCUMENT_TEXT_DETECTION instead of TEXT_DETECTION because prescriptions behave like documents: they may have multiple lines, handwritten sections, labels, instructions, and mixed formatting.",
    )
    add_code(
        doc,
        "POST https://vision.googleapis.com/v1/images:annotate?key=GOOGLE_CLOUD_VISION_API_KEY\n\nfeature.type = DOCUMENT_TEXT_DETECTION\nimage.source.imageUri = Firebase Storage download URL\ntext source = responses[0].fullTextAnnotation.text",
    )

    add_heading(doc, "6. LLM Structure Extraction", 1)
    add_body(
        doc,
        "The raw OCR text is not reliable enough for reminders by itself. The parser service sends it to an LLM with a strict system prompt that requires a JSON array only. The app currently supports OpenAI or Gemini through environment variables.",
    )
    add_code(
        doc,
        'Required JSON fields per item:\n{\n  "medicineName": "string",\n  "dosage": "string",\n  "frequency": "string",\n  "durationDays": number,\n  "instructions": "string"\n}',
    )
    add_body(
        doc,
        "The parser then converts each JSON item into the existing Medication class. This is important because the reminder system expects medication name, dosage, frequency, duration, start/end dates, timings, and instructions in the same shape used by the previous mock data.",
    )

    add_heading(doc, "7. Error Handling", 1)
    add_table(
        doc,
        ["Scenario", "User Message", "App Behavior"],
        [
            [
                "Vision returns no readable text",
                "Could not read prescription. Please retake the photo in good lighting.",
                "The app stays stable and does not navigate with empty OCR text.",
            ],
            [
                "LLM returns invalid JSON or parsing fails",
                "Prescription scanned but medicines could not be extracted. Please add medicines manually.",
                "The app keeps the raw OCR text, uses an empty medication list, and allows manual correction.",
            ],
            [
                "Network/API key/API failure",
                "Unable to scan prescription right now. Please try again.",
                "The app records an error message and leaves state valid.",
            ],
        ],
        widths=[2.1, 2.7, 2.3],
    )

    add_heading(doc, "8. Environment Setup", 1)
    add_body(
        doc,
        "API keys are intentionally not hardcoded. Put real keys into a local .env file. The file is ignored by Git, while .env.example documents the required fields.",
    )
    add_code(
        doc,
        "GOOGLE_CLOUD_VISION_API_KEY=your_google_vision_key\nPRESCRIPTION_LLM_PROVIDER=openai\nOPENAI_API_KEY=your_openai_key\nOPENAI_MODEL=gpt-4o\n\n# Optional Gemini alternative\nGEMINI_API_KEY=your_gemini_key\nGEMINI_MODEL=gemini-1.5-flash",
    )

    add_heading(doc, "9. What Was Not Changed", 1)
    add_bullets(
        doc,
        [
            "Reminder scheduling logic was not modified.",
            "MedicationDoseRepository was not modified.",
            "PrescriptionRepository Firestore write behavior was not modified.",
            "Cloud Functions and FCM reminder logic were not modified.",
            "The OCR output remains compatible because it produces the same Medication list structure that existing downstream code expects.",
        ],
    )

    add_heading(doc, "10. Testing Checklist", 1)
    add_numbered(
        doc,
        [
            "Add valid API keys to .env.",
            "Run flutter pub get.",
            "Open the app and log in as a patient.",
            "Go to Scan Prescription.",
            "Take a clear photo or upload one from gallery.",
            "Confirm the result screen shows raw text and extracted medicines.",
            "Save the prescription.",
            "Open My Medication and confirm future dose reminders were created.",
            "If OCR fails, retake the photo in better lighting.",
            "If medicine extraction fails, manually add or edit medicines before saving.",
        ],
    )

    add_heading(doc, "11. Prescription Test Examples", 1)
    add_body(
        doc,
        "You can print these examples, write them on paper, or display them on another screen and scan them. Start with clean typed text, then test harder cases such as shadows, tilted paper, and handwriting.",
    )

    examples = [
        (
            "Example A - Simple Antibiotic Prescription",
            "Patient: Ali Khan\nDate: 06 May 2026\nRx:\n1. Amoxicillin 500mg - 1 capsule every 8 hours for 7 days. Take after meals.\n2. Paracetamol 500mg - 1 tablet every 6 hours as needed for fever for 3 days. Take with water.\nFollow up if symptoms continue.",
        ),
        (
            "Example B - Twice Daily Medicine",
            "Clinic: MindHeal Family Care\nPatient: Sara Ahmed\nRx:\nCefixime 200mg tablet\nTake 1 tablet twice daily for 5 days after food.\nCetirizine 10mg\nTake 1 tablet at bedtime for 7 days.",
        ),
        (
            "Example C - Diabetes Style Prescription",
            "Patient: Hassan Ali\nDiagnosis: Type 2 Diabetes\nRx:\nMetformin 500mg - 1 tablet twice daily after meals for 30 days.\nGlimepiride 2mg - 1 tablet every morning before breakfast for 30 days.\nCheck blood glucose weekly.",
        ),
        (
            "Example D - Messier Multi-Line Prescription",
            "Rx\nAzithromycin 500 mg\nDay 1: take 1 tablet once daily\ncontinue for 3 days\nIbuprofen 400mg\n1 tab every 8 hours after meal if pain\nORS sachet\nmix 1 sachet in water twice daily for 2 days",
        ),
        (
            "Example E - Minimal Details To Test Defaults",
            "Rx:\nOmeprazole 20mg daily before breakfast\nVitamin D3 once weekly\nLoratadine at night",
        ),
    ]

    for title_text, body_text in examples:
        add_heading(doc, title_text, 2)
        add_code(doc, body_text)

    add_heading(doc, "12. Expected JSON From Parser", 1)
    add_body(
        doc,
        "For Example A, the parser should return a JSON array similar to this. Exact wording can vary, but fields must exist and no value should be null.",
    )
    add_code(
        doc,
        '[\n  {\n    "medicineName": "Amoxicillin",\n    "dosage": "500mg, 1 capsule",\n    "frequency": "every 8 hours",\n    "durationDays": 7,\n    "instructions": "Take after meals"\n  },\n  {\n    "medicineName": "Paracetamol",\n    "dosage": "500mg, 1 tablet",\n    "frequency": "every 6 hours as needed",\n    "durationDays": 3,\n    "instructions": "Take with water"\n  }\n]',
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Prepared for MindHeal OCR implementation and QA")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(110, 110, 110)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
